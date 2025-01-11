import calendar
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
import holidays
from datetime import datetime

# Code generated using ChatGPT

def set_row_height(row, height_cm):
    """Set the height of a Word table row in centimeters."""
    tr = row._tr  # Access the underlying XML element for the row
    trPr = tr.get_or_add_trPr()  # Add the trPr element if it doesn't exist
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 566.93)))  # Convert cm to twips (1 cm = 566.93 twips)
    trHeight.set(qn('w:hRule'), "exact")  # Use "exact" to fix the row height
    trPr.append(trHeight)

def set_cell_padding(cell, padding_cm):
    """Set padding for a table cell in centimeters."""
    tc = cell._tc  # Access the underlying XML element for the cell
    tcPr = tc.get_or_add_tcPr()  # Add the tcPr element if it doesn't exist
    tcMar = OxmlElement('w:tcMar')  # Create a new tcMar (table cell margin) element

    # Define margins (top, bottom, left, right) in twips (1 cm = 566.93 twips)
    margin_value = str(int(padding_cm * 566.93))  # Convert cm to twips
    for margin in ['top', 'bottom', 'left', 'right']:
        element = OxmlElement(f'w:{margin}')
        element.set(qn('w:w'), margin_value)
        element.set(qn('w:type'), 'dxa')
        tcMar.append(element)

    tcPr.append(tcMar)  # Append the margins to the cell properties

def set_cell_color(cell, color):
    """Set the background color of a table cell."""
    cell._element.get_or_add_tcPr()  # Ensure tcPr element exists
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)  # Set the fill color (hex code)
    cell._element.tcPr.append(shading)

def set_red_text(cell):
    """Set the text color of the date number to red for holidays."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)  # RGB for red color

def generate_calendar(year):
    doc = Document()

    # Define "Very Light" pastel colors for each month
    pastel_colors = [
        "DCE9F5",  # Jan Very Light Blue
        "E0F5D3",  # Feb Very Light Green
        "FAD0D4",  # Mar Very Light Pink
        "F9F2D0",  # Apr Very Light Yellow
        "D5E8E5",  # May Very Light Teal
        "F7E5D3",  # Jun Very Light Peach
        "E0F5D3",  # Jul Very Light Green
        "D4E9F7",  # Aug Very Light Sky Blue
        "FAD0D4",  # Sep Very Light Pink
        "F9F2D0",  # Oct Very Light Yellow
        "D5E8E5",  # Nov Very Light Teal
        "e0c2cd",  # Des Very Light Purple
        #"F0E0F8",  # Very Light Lavender
        # "E1F5D6",  # Very Light Mint
        # "F1E1F7",   # Very Light Lilac
        # "F8E3D3",  # Very Light Apricot
    ]

    # Use the holidays package to get Norwegian public holidays for the year
    no_holidays = holidays.Norway(years=year)

    # Calculate total weeks for the year
    total_weeks = sum(len(calendar.monthcalendar(year, month)) for month in range(1, 13))

    # Create a table with enough rows and 7 columns (one per day of the week)
    table = doc.add_table(rows=total_weeks, cols=7)
    table.style = 'Table Grid'

    # Define row height and cell padding
    row_height_cm = 2.25  # Set row height to 2.25 cm
    cell_padding_cm = 0.08  # Set cell padding to 0.08 cm

    # Add calendar days to the table
    row_index = 0
    col_index = calendar.weekday(year, 1, 1)  # Get the weekday of January 1

    for month in range(1, 13):
        month_days = calendar.monthrange(year, month)[1]

        # Set color for the current month
        month_color = pastel_colors[month - 1]

        for day in range(1, month_days + 1):
            cell = table.cell(row_index, col_index)
            cell.text = str(day)

            # Check if the (month, day) is a holiday and set red text color if it is
            date = datetime(year, month, day)
            if date in no_holidays:
                set_red_text(cell)

            # Set font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

            # Set padding for the cell
            set_cell_padding(cell, cell_padding_cm)

            # Set the background color for the cell (Very Light pastel color)
            set_cell_color(cell, month_color)

            # Move to the next cell
            col_index += 1
            if col_index > 6:  # Move to the next row if Sunday is reached
                col_index = 0
                set_row_height(table.rows[row_index], row_height_cm)  # Set row height
                row_index += 1

    # Remove any extra empty rows at the end
    for row in table.rows:
        if all(cell.text == '' for cell in row.cells):
            tbl = table._element
            tbl.remove(row._element)

    # Ensure the last row has the correct height
    if row_index < len(table.rows):
        set_row_height(table.rows[row_index], row_height_cm)

    # Save the document
    file_name = f'{year}_Calendar.docx'
    doc.save(file_name)
    print(f"Calendar for {year} saved as '{file_name}'")

# Example usage
generate_calendar(2025)
